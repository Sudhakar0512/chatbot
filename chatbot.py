#only PDF

# import streamlit as st
# from PyPDF2 import PdfReader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)


# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text


# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks


# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")


# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """

#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain


# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload PDF files and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")


# # Main app with enhanced UI
# def main():
#     st.set_page_config(page_title="PDF Query Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         .sidebar {
#             background-color: #f0f0f0;
#         }
#         .menu-title {
#             font-size: 22px;
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'> PDF Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area with modern styling
#     user_question = st.text_input("Ask a Question from the PDF Files", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for PDF upload
#     with st.sidebar:
#         st.markdown("<div class='menu-title'>Menu</div>", unsafe_allow_html=True)
#         pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True, type="pdf")
        
#         # Process button with modern styling
#         if st.button("Submit & Process", key="submit-btn"):
#             if pdf_docs:
#                 with st.spinner("Processing..."):
#                     try:
#                         raw_text = get_pdf_text(pdf_docs)
#                         if raw_text.strip() == "":
#                             st.error("The uploaded PDFs do not contain any readable text.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing PDFs.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload PDF files before clicking Submit.")


# if __name__ == "__main__":
#     main()

# # PDF and Image
# import streamlit as st
# from PyPDF2 import PdfReader
# import pytesseract
# from PIL import Image
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)


# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text


# def get_image_text(image_files):
#     text = ""
#     for image in image_files:
#         img = Image.open(image)
#         extracted_text = pytesseract.image_to_string(img)
#         text += extracted_text
#     return text


# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks


# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")


# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """

#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain


# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload PDF and image files and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")


# # Main app with enhanced UI
# def main():
#     st.set_page_config(page_title=" Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         .sidebar {
#             background-color: #f0f0f0;
#         }
#         .menu-title {
#             font-size: 22px;
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'> PDF & Image Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area with modern styling
#     user_question = st.text_input("Ask a Question from the PDF and Image Files", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for PDF and image upload
#     with st.sidebar:
#         st.markdown("<div class='menu-title'>File Formates</div>", unsafe_allow_html=True)
#         pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True, type="pdf")
#         image_files = st.file_uploader("Upload your Image Files", accept_multiple_files=True, type=["jpg", "jpeg", "png"])
        
#         # Process button with modern styling
#         if st.button("Submit & Process", key="submit-btn"):
#             if pdf_docs or image_files:
#                 with st.spinner("Processing..."):
#                     try:
#                         pdf_text = get_pdf_text(pdf_docs) if pdf_docs else ""
#                         image_text = get_image_text(image_files) if image_files else ""
#                         raw_text = pdf_text + image_text
                        
#                         if raw_text.strip() == "":
#                             st.error("The uploaded files do not contain any readable text.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing files.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload PDF or image files before clicking Submit.")


# if __name__ == "__main__":
#     main()

# PDF Image CSV
# import streamlit as st
# from PyPDF2 import PdfReader
# import pytesseract
# from PIL import Image
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv
# import pandas as pd  # New import for CSV handling

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)


# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text


# def get_image_text(image_files):
#     text = ""
#     for image in image_files:
#         img = Image.open(image)
#         extracted_text = pytesseract.image_to_string(img)
#         text += extracted_text
#     return text


# # New function to extract text from CSV files
# def get_csv_text(csv_files):
#     text = ""
#     for csv_file in csv_files:
#         try:
#             df = pd.read_csv(csv_file)
#             # Convert the DataFrame into a string for processing
#             text += df.to_string(index=False)
#         except Exception as e:
#             st.warning(f"Error reading {csv_file.name}: {str(e)}")
#     return text


# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks


# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")


# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """

#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain


# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload PDF, image, or CSV files and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")


# # Main app with enhanced UI
# def main():
#     st.set_page_config(page_title=" Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         .sidebar {
#             background-color: #f0f0f0;
#         }
#         .menu-title {
#             font-size: 22px;
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'> PDF, Image & CSV Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area with modern styling
#     user_question = st.text_input("Ask a Question from the PDF, Image, and CSV Files", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for PDF, image, and CSV upload
#     with st.sidebar:
#         st.markdown("<div class='menu-title'>File Formates</div>", unsafe_allow_html=True)
#         pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True, type="pdf")
#         image_files = st.file_uploader("Upload your Image Files", accept_multiple_files=True, type=["jpg", "jpeg", "png"])
#         csv_files = st.file_uploader("Upload your CSV Files", accept_multiple_files=True, type="csv")  # New CSV uploader

#         # Process button with modern styling
#         if st.button("Submit & Process", key="submit-btn"):
#             if pdf_docs or image_files or csv_files:
#                 with st.spinner("Processing..."):
#                     try:
#                         pdf_text = get_pdf_text(pdf_docs) if pdf_docs else ""
#                         image_text = get_image_text(image_files) if image_files else ""
#                         csv_text = get_csv_text(csv_files) if csv_files else ""  # New CSV processing
#                         raw_text = pdf_text + image_text + csv_text
                        
#                         if raw_text.strip() == "":
#                             st.error("The uploaded files do not contain any readable text.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing files.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload PDF, image, or CSV files before clicking Submit.")


# if __name__ == "__main__":
#     main()


#PDF Image CSV Word
# import streamlit as st
# from PyPDF2 import PdfReader
# import pytesseract
# from PIL import Image
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv
# import pandas as pd  # For CSV handling
# import docx  # New import for Word file handling

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)


# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text


# def get_image_text(image_files):
#     text = ""
#     for image in image_files:
#         img = Image.open(image)
#         extracted_text = pytesseract.image_to_string(img)
#         text += extracted_text
#     return text


# # New function to extract text from CSV files
# def get_csv_text(csv_files):
#     text = ""
#     for csv_file in csv_files:
#         try:
#             df = pd.read_csv(csv_file)
#             # Convert the DataFrame into a string for processing
#             text += df.to_string(index=False)
#         except Exception as e:
#             st.warning(f"Error reading {csv_file.name}: {str(e)}")
#     return text


# # New function to extract text from Word files
# def get_word_text(word_files):
#     text = ""
#     for word_file in word_files:
#         try:
#             doc = docx.Document(word_file)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         except Exception as e:
#             st.warning(f"Error reading {word_file.name}: {str(e)}")
#     return text


# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks


# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")


# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """

#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain


# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload PDF, image, CSV, or Word files and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")


# # Main app with enhanced UI
# def main():
#     st.set_page_config(page_title=" Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         .sidebar {
#             background-color: #f0f0f0;
#         }
#         .menu-title {
#             font-size: 22px;
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'> PDF, Image, CSV & Word Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area with modern styling
#     user_question = st.text_input("Ask a Question from the PDF, Image, CSV, and Word Files", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for PDF, image, CSV, and Word upload
#     with st.sidebar:
#         st.markdown("<div class='menu-title'>File Formates</div>", unsafe_allow_html=True)
#         pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True, type="pdf")
#         image_files = st.file_uploader("Upload your Image Files", accept_multiple_files=True, type=["jpg", "jpeg", "png"])
#         csv_files = st.file_uploader("Upload your CSV Files", accept_multiple_files=True, type="csv")
#         word_files = st.file_uploader("Upload your Word Files", accept_multiple_files=True, type="docx")  # New Word uploader

#         # Process button with modern styling
#         if st.button("Submit & Process", key="submit-btn"):
#             if pdf_docs or image_files or csv_files or word_files:
#                 with st.spinner("Processing..."):
#                     try:
#                         pdf_text = get_pdf_text(pdf_docs) if pdf_docs else ""
#                         image_text = get_image_text(image_files) if image_files else ""
#                         csv_text = get_csv_text(csv_files) if csv_files else ""  # CSV processing
#                         word_text = get_word_text(word_files) if word_files else ""  # New Word processing
#                         raw_text = pdf_text + image_text + csv_text + word_text
                        
#                         if raw_text.strip() == "":
#                             st.error("The uploaded files do not contain any readable text.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing files.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload PDF, image, CSV, or Word files before clicking Submit.")


# if __name__ == "__main__":
#     main()



# all formates
# import streamlit as st
# from PyPDF2 import PdfReader
# import pytesseract
# from PIL import Image
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv
# import pandas as pd  # For CSV handling
# import docx  # Word file handling

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)

# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text

# def get_image_text(image_files):
#     text = ""
#     for image in image_files:
#         img = Image.open(image)
#         extracted_text = pytesseract.image_to_string(img)
#         text += extracted_text
#     return text

# def get_csv_text(csv_files):
#     text = ""
#     for csv_file in csv_files:
#         try:
#             df = pd.read_csv(csv_file)
#             text += df.to_string(index=False)
#         except Exception as e:
#             st.warning(f"Error reading {csv_file.name}: {str(e)}")
#     return text

# def get_word_text(word_files):
#     text = ""
#     for word_file in word_files:
#         try:
#             doc = docx.Document(word_file)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         except Exception as e:
#             st.warning(f"Error reading {word_file.name}: {str(e)}")
#     return text

# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks

# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")

# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """
#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain

# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload PDF, image, CSV, Word files or input text and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")

# # Main app with enhanced UI
# def main():
#     st.set_page_config(page_title=" Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         .sidebar {
#             background-color: #f0f0f0;
#         }
#         .menu-title {
#             font-size: 22px;
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'> PDF, Image, CSV, Word, and Text Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area with modern styling
#     user_question = st.text_input("Ask a Question from the PDF, Image, CSV, Word, and Text Files", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for file and text input
#     with st.sidebar:
#         st.markdown("<div class='menu-title'>File Formats</div>", unsafe_allow_html=True)
#         pdf_docs = st.file_uploader("Upload your PDF Files", accept_multiple_files=True, type="pdf")
#         image_files = st.file_uploader("Upload your Image Files", accept_multiple_files=True, type=["jpg", "jpeg", "png"])
#         csv_files = st.file_uploader("Upload your CSV Files", accept_multiple_files=True, type="csv")
#         word_files = st.file_uploader("Upload your Word Files", accept_multiple_files=True, type="docx")
#         text_input = st.text_area("Enter plain text", help="You can directly input plain text here.")  # New text input

#         # Process button with modern styling
#         if st.button("Submit & Process", key="submit-btn"):
#             if pdf_docs or image_files or csv_files or word_files or text_input:
#                 with st.spinner("Processing..."):
#                     try:
#                         pdf_text = get_pdf_text(pdf_docs) if pdf_docs else ""
#                         image_text = get_image_text(image_files) if image_files else ""
#                         csv_text = get_csv_text(csv_files) if csv_files else ""
#                         word_text = get_word_text(word_files) if word_files else ""
#                         raw_text = pdf_text + image_text + csv_text + word_text + text_input  # Add text input

#                         if raw_text.strip() == "":
#                             st.error("The uploaded files or input text do not contain any readable content.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing files and text.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload files or enter text before clicking Submit.")

# if __name__ == "__main__":
#     main()













# import streamlit as st
# from PyPDF2 import PdfReader
# import pytesseract
# from PIL import Image
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# import google.generativeai as genai
# from langchain.vectorstores import FAISS
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv
# import pandas as pd  # For CSV handling
# import docx  # Word file handling

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("Google API key not found. Please check the environment variables.")
# else:
#     genai.configure(api_key=api_key)

# # Functions to extract text from different file formats
# def get_pdf_text(pdf_files):
#     text = ""
#     for pdf in pdf_files:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text
#             else:
#                 st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
#     return text

# def get_image_text(image_files):
#     text = ""
#     for image in image_files:
#         img = Image.open(image)
#         extracted_text = pytesseract.image_to_string(img)
#         text += extracted_text
#     return text

# def get_csv_text(csv_files):
#     text = ""
#     for csv_file in csv_files:
#         try:
#             df = pd.read_csv(csv_file)
#             text += df.to_string(index=False)
#         except Exception as e:
#             st.warning(f"Error reading {csv_file.name}: {str(e)}")
#     return text

# def get_word_text(word_files):
#     text = ""
#     for word_file in word_files:
#         try:
#             doc = docx.Document(word_file)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         except Exception as e:
#             st.warning(f"Error reading {word_file.name}: {str(e)}")
#     return text

# def get_text_chunks(text):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#     chunks = text_splitter.split_text(text)
#     return chunks

# def get_vector_store(text_chunks):
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#     except Exception as e:
#         st.error(f"Failed to create vector store: {str(e)}")

# def get_conversational_chain():
#     prompt_template = """
#     Answer the question as detailed as possible from the provided context. If the answer is not in 
#     the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
#     Context:\n{context}?\n
#     Question:\n{question}\n

#     Answer:
#     """
#     model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#     prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#     chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#     return chain

# def user_input(user_question):
#     embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

#     if not os.path.exists("faiss_index/index.faiss"):
#         st.error("FAISS index not found. Please upload files or input text and process them first.")
#         return

#     try:
#         new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#         docs = new_db.similarity_search(user_question)
#         chain = get_conversational_chain()

#         response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
#         st.write("Reply:", response["output_text"])
#     except Exception as e:
#         st.error(f"Error during user input processing: {str(e)}")

# # Main app with single file uploader and text input
# def main():
#     st.set_page_config(page_title="Chatbot", layout="centered")

#     # Centered chat window-style UI
#     st.markdown("""
#         <style>
#         .chatbox {
#             background-color: #f9f9f9;
#             padding: 10px;
#             border-radius: 10px;
#             border: 1px solid #ddd;
#         }
#         .input-text {
#             width: 100%;
#             border-radius: 8px;
#             padding: 10px;
#             border: 1px solid #ddd;
#             font-size: 16px;
#         }
#         .submit-btn {
#             background-color: #0088cc;
#             color: white;
#             border-radius: 8px;
#             padding: 10px 20px;
#             font-size: 16px;
#             border: none;
#             cursor: pointer;
#         }
#         .submit-btn:hover {
#             background-color: #0077b3;
#         }
#         </style>
#         """, unsafe_allow_html=True)

#     st.markdown("<h2 style='text-align:center;'>PDF, Image, CSV, Word & Text Query Chatbot</h2>", unsafe_allow_html=True)

#     # Chat input area
#     user_question = st.text_input("Ask a Question from the uploaded files and text", key="input-text", help="Type your query here...")

#     if user_question:
#         user_input(user_question)

#     # Sidebar for single file upload and text input
#     with st.sidebar:
#         st.markdown("<div style='font-size: 18px; font-weight: bold;'>Upload Files & Input Text</div>", unsafe_allow_html=True)
        
#         # Single file uploader for multiple formats
#         files = st.file_uploader("Upload files (PDF, CSV, Word, Image)", accept_multiple_files=True, type=["pdf", "csv", "docx", "jpg", "jpeg", "png"])

#         # Text input field for manual input
#         text_input = st.text_area("Or input plain text", help="You can directly input plain text here.")

#         # Process button
#         if st.button("Submit & Process"):
#             if files or text_input:
#                 with st.spinner("Processing..."):
#                     try:
#                         # Process uploaded files based on type
#                         pdf_text = get_pdf_text([file for file in files if file.name.endswith(".pdf")])
#                         image_text = get_image_text([file for file in files if file.name.endswith((".jpg", ".jpeg", ".png"))])
#                         csv_text = get_csv_text([file for file in files if file.name.endswith(".csv")])
#                         word_text = get_word_text([file for file in files if file.name.endswith(".docx")])
                        
#                         # Combine all text inputs
#                         raw_text = pdf_text + image_text + csv_text + word_text + text_input
                        
#                         if raw_text.strip() == "":
#                             st.error("The uploaded files or input text do not contain any readable content.")
#                         else:
#                             text_chunks = get_text_chunks(raw_text)
#                             get_vector_store(text_chunks)
#                             st.success("Done processing files and text.")
#                     except Exception as e:
#                         st.error(f"Error during processing: {str(e)}")
#             else:
#                 st.error("Please upload files or enter text before clicking Submit.")

# if __name__ == "__main__":
#     main()


# ## FAISS update
import streamlit as st
from PyPDF2 import PdfReader
import pytesseract
from PIL import Image
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
# from langchain.vectorstores import FAISS
from langchain_community.vectorstores import FAISS

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import pandas as pd  # For CSV handling
import docx  # Word file handling

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("Google API key not found. Please check the environment variables.")
else:
    genai.configure(api_key=api_key)

# Functions to extract text from different file formats
def get_pdf_text(pdf_files):
    text = ""
    for pdf in pdf_files:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text
            else:
                st.warning(f"Page {page} in {pdf.name} does not contain readable text.")
    return text

def get_image_text(image_files):
    text = ""
    for image in image_files:
        img = Image.open(image)
        extracted_text = pytesseract.image_to_string(img)
        text += extracted_text
    return text

def get_csv_text(csv_files):
    text = ""
    for csv_file in csv_files:
        try:
            df = pd.read_csv(csv_file)
            text += df.to_string(index=False)
        except Exception as e:
            st.warning(f"Error reading {csv_file.name}: {str(e)}")
    return text

def get_word_text(word_files):
    text = ""
    for word_file in word_files:
        try:
            doc = docx.Document(word_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            st.warning(f"Error reading {word_file.name}: {str(e)}")
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def update_vector_store(new_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    
    # Load existing FAISS index
    if os.path.exists("faiss_index/index.faiss"):
        vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        vector_store.add_texts(new_chunks)  # Add new chunks to the existing index
    else:
        vector_store = FAISS.from_texts(new_chunks, embedding=embeddings)  # Create new index if it doesn't exist
    
    vector_store.save_local("faiss_index")  # Save the updated index

def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context. If the answer is not in 
    the provided context, just say, "answer is not available in the context," and don't provide the wrong answer.\n\n
    Context:\n{context}?\n
    Question:\n{question}\n

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain

def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

    if not os.path.exists("faiss_index/index.faiss"):
        st.error("FAISS index not found. Please upload files or input text and process them first.")
        return

    try:
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)
        chain = get_conversational_chain()

        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        st.write("Reply:", response["output_text"])
    except Exception as e:
        st.error(f"Error during user input processing: {str(e)}")

# Main app with single file uploader and text input
def main():
    st.set_page_config(page_title="Chatbot", layout="centered")

    # Centered chat window-style UI
    st.markdown("""
        <style>
        .chatbox {
            background-color: #f9f9f9;
            padding: 10px;
            border-radius: 10px;
            border: 1px solid #ddd;
        }
        .input-text {
            width: 100%;
            border-radius: 8px;
            padding: 10px;
            border: 1px solid #ddd;
            font-size: 16px;
        }
        .submit-btn {
            background-color: #0088cc;
            color: white;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 16px;
            border: none;
            cursor: pointer;
        }
        .submit-btn:hover {
            background-color: #0077b3;
        }
        </style>
        """, unsafe_allow_html=True)

    st.markdown("<h2 style='text-align:center;'>PDF, Image, CSV, Word & Text Query Chatbot</h2>", unsafe_allow_html=True)

    # Chat input area
    user_question = st.text_input("Ask a Question from the uploaded files and text", key="input-text", help="Type your query here...")

    if user_question:
        user_input(user_question)

    # Sidebar for single file upload and text input
    with st.sidebar:
        st.markdown("<div style='font-size: 18px; font-weight: bold;'>Upload Files & Input Text</div>", unsafe_allow_html=True)
        
        # Single file uploader for multiple formats
        files = st.file_uploader("Upload files (PDF, CSV, Word, Image)", accept_multiple_files=True, type=["pdf", "csv", "docx", "jpg", "jpeg", "png"])

        # Text input field for manual input
        text_input = st.text_area("Or input plain text", help="You can directly input plain text here.")

        # Process button
        if st.button("Submit & Process"):
            if files or text_input:
                with st.spinner("Processing..."):
                    try:
                        # Process uploaded files based on type
                        pdf_text = get_pdf_text([file for file in files if file.name.endswith(".pdf")])
                        image_text = get_image_text([file for file in files if file.name.endswith((".jpg", ".jpeg", ".png"))])
                        csv_text = get_csv_text([file for file in files if file.name.endswith(".csv")])
                        word_text = get_word_text([file for file in files if file.name.endswith(".docx")])
                        
                        # Combine all text inputs
                        raw_text = pdf_text + image_text + csv_text + word_text + text_input
                        
                        if raw_text.strip() == "":
                            st.error("The uploaded files or input text do not contain any readable content.")
                        else:
                            text_chunks = get_text_chunks(raw_text)
                            update_vector_store(text_chunks)  # Update vector store with new data
                            st.success("Done processing files and text.")
                    except Exception as e:
                        st.error(f"Error during processing: {str(e)}")
            else:
                st.error("Please upload files or enter text before clicking Submit.")

if __name__ == "__main__":
    main()
